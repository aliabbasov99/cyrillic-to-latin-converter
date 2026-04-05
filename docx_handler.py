"""
DOCX processing module.

Uses python-docx to:
  1. Open a .docx file
  2. Walk every paragraph → every run
  3. Transliterate Cyrillic → Latin using converter.convert_text()
  4. Preserve all formatting (bold, italic, font name, font size, color, etc.)
  5. Also processes tables (each cell's paragraphs)
  6. Also processes headers and footers
  7. Save as a new .docx
"""

import os
from docx import Document
from converter import convert_text


def process_docx(input_path: str, output_path=None, replace_quotes: bool = False) -> str:
    """
    Convert all Cyrillic text in a DOCX to Latin script.

    Args:
        input_path:  Path to the source .docx file.
        output_path: Where to save the result. If None, auto-generates
                     a path like ``original_latin.docx``.
        replace_quotes: If True, replaces « and » with ".

    Returns:
        The absolute path of the generated output file.
    """
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_latin{ext}"

    doc = Document(input_path)

    # ── Main body paragraphs ──
    for para in doc.paragraphs:
        _convert_paragraph(para, replace_quotes)

    # ── Tables ──
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _convert_paragraph(para, replace_quotes)

    # ── Headers & Footers ──
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    _convert_paragraph(para, replace_quotes)

        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    _convert_paragraph(para, replace_quotes)

    doc.save(output_path)

    return os.path.abspath(output_path)


def _convert_paragraph(para, replace_quotes: bool = False) -> None:
    """Convert text in every run of a paragraph, preserving formatting."""
    for run in para.runs:
        original = run.text
        converted = convert_text(original, replace_quotes)
        if converted != original:
            run.text = converted
